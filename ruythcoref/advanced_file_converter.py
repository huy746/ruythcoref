import os
import json
import pandas as pd
from shutil import copyfile
from docx import Document
import PyPDF2
import xml.etree.ElementTree as ET

class AdvancedFileConverter:
    @staticmethod
    def copy_file(src_path, dest_path):
        if not os.path.exists(src_path):
            raise FileNotFoundError(f"File {src_path} không tồn tại.")
        copyfile(src_path, dest_path)
        print(f"File đã được sao chép sang {dest_path}")

    @staticmethod
    def read_txt(file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            return [line.strip() for line in f.readlines()]

    @staticmethod
    def write_txt(file_path, lines):
        with open(file_path, 'w', encoding='utf-8') as f:
            for line in lines:
                f.write(line + '\n')

    @staticmethod
    def read_docx(file_path):
        doc = Document(file_path)
        return [p.text for p in doc.paragraphs if p.text.strip() != '']

    @staticmethod
    def write_docx(file_path, lines):
        doc = Document()
        for line in lines:
            doc.add_paragraph(line)
        doc.save(file_path)

    @staticmethod
    def read_pdf(file_path):
        reader = PyPDF2.PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text())
        return text

    @staticmethod
    def write_pdf(file_path, lines):
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        for line in lines:
            pdf.multi_cell(0, 10, line)
        pdf.output(file_path)

    @staticmethod
    def read_csv(file_path):
        return pd.read_csv(file_path)

    @staticmethod
    def write_csv(file_path, df):
        df.to_csv(file_path, index=False)

    @staticmethod
    def read_json(file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    @staticmethod
    def write_json(file_path, data):
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4)

    @staticmethod
    def read_xml(file_path):
        tree = ET.parse(file_path)
        return tree.getroot()

    @staticmethod
    def write_xml(file_path, root):
        tree = ET.ElementTree(root)
        tree.write(file_path, encoding='utf-8', xml_declaration=True)

    @staticmethod
    def convert(src_path, dest_path):
        if not os.path.exists(src_path):
            raise FileNotFoundError(f"File {src_path} không tồn tại.")

        src_ext = os.path.splitext(src_path)[1].lower()
        dest_ext = os.path.splitext(dest_path)[1].lower()

        # Đọc nội dung
        if src_ext == '.txt':
            content = AdvancedFileConverter.read_txt(src_path)
        elif src_ext == '.docx':
            content = AdvancedFileConverter.read_docx(src_path)
        elif src_ext == '.pdf':
            content = AdvancedFileConverter.read_pdf(src_path)
        elif src_ext == '.csv':
            content = AdvancedFileConverter.read_csv(src_path)
        elif src_ext == '.json':
            content = AdvancedFileConverter.read_json(src_path)
        elif src_ext == '.xml':
            content = AdvancedFileConverter.read_xml(src_path)
        else:
            raise ValueError(f"Định dạng {src_ext} không được hỗ trợ.")

        # Ghi ra file đích
        if dest_ext == '.txt':
            if isinstance(content, pd.DataFrame):
                content = content.to_csv(index=False).splitlines()
            elif isinstance(content, str):
                content = [content]
            AdvancedFileConverter.write_txt(dest_path, content)
        elif dest_ext == '.docx':
            if not isinstance(content, list):
                content = [str(content)]
            AdvancedFileConverter.write_docx(dest_path, content)
        elif dest_ext == '.pdf':
            if not isinstance(content, list):
                content = [str(content)]
            AdvancedFileConverter.write_pdf(dest_path, content)
        elif dest_ext == '.csv':
            if isinstance(content, pd.DataFrame):
                AdvancedFileConverter.write_csv(dest_path, content)
            else:
                AdvancedFileConverter.write_csv(dest_path, pd.DataFrame({"text": content}))
        elif dest_ext == '.json':
            if isinstance(content, pd.DataFrame):
                content = content.to_dict(orient='records')
            AdvancedFileConverter.write_json(dest_path, content)
        elif dest_ext == '.xml':
            if not isinstance(content, ET.Element):
                root = ET.Element("root")
                for i, item in enumerate(content):
                    e = ET.SubElement(root, f"item{i}")
                    e.text = str(item)
                content = root
            AdvancedFileConverter.write_xml(dest_path, content)
        else:
            raise ValueError(f"Định dạng {dest_ext} không được hỗ trợ.")

        print(f"Chuyển đổi {src_path} → {dest_path} thành công!")
  
